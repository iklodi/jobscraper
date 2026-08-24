import db
import os
import json
import datetime
from docx import Document
from docx.shared import Pt
import subprocess
from groq import Groq
from google import genai
from google.genai import types
import time
import progress_tracker
from playwright.async_api import async_playwright
import PyPDF2

CVS_DIR = os.environ.get('CVS_DIR', 'cvs')

def get_pdf_page_count(pdf_path):
    try:
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return len(reader.pages)
    except Exception as e:
        print(f"Error reading PDF page count for {pdf_path}: {e}")
        return None

def load_prompt(filename):
    path = os.path.join(os.path.dirname(__file__), 'prompts', filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

# Models
CV_TEMPLATE_NAME = os.environ.get('CV_TEMPLATE_NAME', 'Base_CV_Template.docx')
DOSSIER_NAME = os.environ.get('DOSSIER_NAME', 'Career_Dossier.md')
CL_TEMPLATE_NAME = os.environ.get('CL_TEMPLATE_NAME', 'Base_CL_Template.docx')

CV_PATH = os.path.join(CVS_DIR, 'docs', CV_TEMPLATE_NAME)
DOSSIER_PATH = os.path.join(CVS_DIR, 'docs', DOSSIER_NAME)
CL_PATH = os.path.join(CVS_DIR, 'docs', CL_TEMPLATE_NAME)
OUTPUT_DIR = os.path.join(CVS_DIR, 'applications')
GROQ_MODEL = 'llama-3.3-70b-versatile'
GEMINI_MODEL = 'gemini-3.5-flash'

def get_groq_client():
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return None
    return Groq(api_key=api_key)

def get_gemini_client():
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return None
    return genai.Client(api_key=api_key)

def generate_tailored_texts(groq_client, gemini_client, job, cv_text, dossier_text, custom_instructions=None):
    job_id, title, company, location, description, link, score, reasoning, status, created_at, is_promoted, estimated_salary, is_recruiter, hiring_manager_name, jd_language, application_notes = job
    
    address_name = hiring_manager_name.split()[0] if hiring_manager_name else "the hiring manager"
    language_instruction = f"- Language Instruction: The job description is in {jd_language}. YOU MUST WRITE THE COVER LETTER ENTIRELY IN {jd_language}." if jd_language else ""
    
    candidate_name = os.environ.get('CANDIDATE_NAME', 'Jane Doe')
    
    custom_block = f"\nCRITICAL CUSTOM INSTRUCTIONS FROM THE USER (these override any conflicting rules below, including Task 5):\n{custom_instructions}\n" if custom_instructions else ""

    prompt_template = load_prompt('generator_prompt.txt')
    prompt = prompt_template.replace('{custom_instructions}', custom_block) \
                            .replace('{title}', title) \
                            .replace('{company}', company) \
                            .replace('{description}', description) \
                            .replace('{cv_text}', cv_text) \
                            .replace('{dossier_text}', dossier_text) \
                            .replace('{address_name}', address_name) \
                            .replace('{candidate_name}', candidate_name) \
                            .replace('{language_instruction}', language_instruction)
        
    max_retries = 5
    for attempt in range(max_retries):
        result = None
        error_msg = ""
        
        if gemini_client:
            try:
                response = gemini_client.models.generate_content(
                    model=GEMINI_MODEL,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                    ),
                )
                import json
                result = json.loads(response.text)
            except Exception as e:
                error_msg += f"Gemini Error: {str(e)} | "
                
        if not result and groq_client:
            try:
                response = groq_client.chat.completions.create(
                    model=GROQ_MODEL,
                    messages=[{"role": "user", "content": prompt}],
                    response_format={"type": "json_object"},
                    temperature=0.1
                )
                import json
                result = json.loads(response.choices[0].message.content)
            except Exception as e:
                error_msg += f"Groq Error: {str(e)}"
                
        if result:
            break
            
        if '429' in error_msg or 'RESOURCE_EXHAUSTED' in error_msg or 'rate_limit' in error_msg:
            print(f"API Rate Limit Block: {error_msg}")
            print("Falling back to 'Slow and Smart' mode (sleeping 32 seconds to clear TPM quota)...")
            import time
            time.sleep(32)
        elif '503' in error_msg or 'UNAVAILABLE' in error_msg:
            print(f"API overload detected: {error_msg}")
            import time
            time.sleep(10)
        else:
            raise Exception(f"Failed to generate documents from APIs: {error_msg}")
    
    # Aggressive post-processing fallback just in case the AI ignores the prompt
    result['cv_summary'] = result['cv_summary'].replace('—', ' - ').replace('–', ' - ')
    result['cover_letter_body'] = result['cover_letter_body'].replace('—', ' - ').replace('–', ' - ')
    return result

def extract_text(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def adapt_cv(base_cv_path, new_cv_path, texts):
    doc = Document(base_cv_path)
    summary_text = texts.get('cv_summary', '')
    replacements = texts.get('cv_replacements', [])
    removals = texts.get('cv_removals', [])
    
    replaced_summary = False
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
            
        # 1. Replace Summary
        if len(p_text) > 100 and not replaced_summary:
            if p.runs:
                p.runs[0].text = summary_text
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.text = summary_text
            replaced_summary = True
            continue
            
        # 2. Check Removals
        if any(rem.strip() and rem.strip() in p_text for rem in removals):
            delete_paragraph(p)
            continue
            
        # 3. Check Replacements
        for rep in list(replacements):
            old_val = rep.get('old', '').strip()
            new_val = rep.get('new', '')
            if old_val and new_val and old_val in p_text:
                import re
                # Use regex with word boundaries to prevent partial word replacements like "Architect" replacing inside "Architecture"
                # If old_val contains special characters, re.escape handles them.
                pattern = r'\b' + re.escape(old_val) + r'\b'
                if re.search(pattern, p.text):
                    full_text = re.sub(pattern, new_val, p.text)
                    font_name = "Verdana"
                    font_size = None
                    if p.runs:
                        if p.runs[0].font.name:
                            font_name = p.runs[0].font.name
                        if p.runs[0].font.size:
                            font_size = p.runs[0].font.size
                            
                    p.clear()
                    run = p.add_run(full_text)
                    run.font.name = font_name
                    if font_size:
                        run.font.size = font_size
                    
                    replacements.remove(rep)
                    break
                elif old_val == p_text:
                    # Fallback for exact matches where word boundaries might fail (e.g. symbols)
                    replace_text_in_paragraph(p, old_val, new_val)
                    replacements.remove(rep)
                    break
                
    doc.save(new_cv_path)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        full_text = paragraph.text.replace(old_text, new_text)
        font_name = "Verdana"
        font_size = None
        if paragraph.runs:
            if paragraph.runs[0].font.name:
                font_name = paragraph.runs[0].font.name
            if paragraph.runs[0].font.size:
                font_size = paragraph.runs[0].font.size
                
        paragraph.clear()
        run = paragraph.add_run(full_text)
        run.font.name = font_name
        if font_size:
            run.font.size = font_size

def sanitize_generated_texts(texts, company, location, jd_language):
    """Deterministic guards for AI output fields that end up verbatim in the documents."""
    import re

    # Header company: the posting company beats a generic punt. If it's an agency,
    # "<Agency> Hiring Team" is still a correct addressee.
    dc = (texts.get('display_company') or '').strip()
    if not dc or dc.lower() in ('hiring team', 'unknown', 'n/a', 'none'):
        texts['display_company'] = company

    # HQ: drop 'Unknown'-style fragments; fall back to the job's location
    hq = (texts.get('company_hq') or '').strip()
    parts = [p.strip() for p in hq.split(',')
             if p.strip() and p.strip().lower() not in ('unknown', 'n/a', 'none')]
    texts['company_hq'] = ', '.join(parts) if parts else re.sub(r'\s*\([^)]*\)', '', location or '').strip()

    # Greeting: the first line must be a well-formed salutation, never a bare name
    lines = (texts.get('cover_letter_body') or '').split('\n')
    starters = ('dear', 'hi', 'hello', 'greetings', 'bonjour', 'cher', 'chère',
                'madame', 'monsieur', 'sehr geehrte', 'guten tag', 'hallo', 'liebe')
    for i, line in enumerate(lines):
        first = line.strip()
        if not first:
            continue
        if any(first.lower().startswith(s) for s in starters):
            if first[-1] not in ',:':
                lines[i] = f"{first},"
        elif len(first) <= 40 and len(first.split()) <= 4 and first[-1] not in '.!?:':
            jl = (jd_language or '').lower()
            prefix = 'Bonjour' if 'french' in jl else 'Guten Tag' if 'german' in jl else 'Dear'
            lines[i] = f"{prefix} {first.rstrip(',')},"
        break
    texts['cover_letter_body'] = '\n'.join(lines)
    return texts

def adapt_cl(base_cl_path, new_cl_path, body_text, company, display_company, location, hiring_manager_name, jd_language):
    doc = Document(base_cl_path)
    
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower().startswith("dear "):
            start_idx = i
            break
            
    if start_idx != -1:
        import re
        company_location_clean = re.sub(r'\s*\([^)]*\)', '', location).strip()
        
        # Replace the hard-coded addressee info in the header (paragraphs before "Dear ")
        today_str = datetime.datetime.now().strftime("%d.%m.%Y")
        for p in doc.paragraphs[:start_idx]:
            if jd_language and "french" in jd_language.lower():
                # Localise the sender's own city line for French-language letters.
                home_city = os.environ.get('HOME_CITY', '')
                home_city_fr = os.environ.get('HOME_CITY_FR', '')
                if home_city and home_city_fr:
                    replace_text_in_paragraph(p, home_city, home_city_fr)
                
            if display_company and display_company.lower() != "hiring team":
                replace_text_in_paragraph(p, "[COMPANY]", display_company)
            else:
                # Template already reads "[COMPANY] Hiring Team" — drop the
                # placeholder (and its trailing space) instead of injecting
                # another "Hiring Team" in front of it.
                replace_text_in_paragraph(p, "[COMPANY] ", "")
                replace_text_in_paragraph(p, "[COMPANY]", "")
                
            replace_text_in_paragraph(p, "[DATE]", today_str)
            
            if location:
                replace_text_in_paragraph(p, "[LOCATION]", company_location_clean)
            else:
                replace_text_in_paragraph(p, "\n[LOCATION]", "")
                replace_text_in_paragraph(p, "[LOCATION]\n", "")
                replace_text_in_paragraph(p, "[LOCATION]", "")

        # Find the reference paragraph for styling (the first actual body paragraph)
        ref_p = doc.paragraphs[start_idx]
        for i in range(start_idx, len(doc.paragraphs)):
            if len(doc.paragraphs[i].text) > 50:
                ref_p = doc.paragraphs[i]
                break
                
        style = ref_p.style
        alignment = ref_p.alignment
        font_name = "Verdana"
        font_size = None
        if ref_p.runs:
            if ref_p.runs[0].font.name:
                font_name = ref_p.runs[0].font.name
            if ref_p.runs[0].font.size:
                font_size = ref_p.runs[0].font.size
        
        # Delete old letter body
        paragraphs_to_delete = doc.paragraphs[start_idx:]
        for p in paragraphs_to_delete:
            delete_paragraph(p)
            
        # Append new tailored letter body
        for text_block in body_text.split('\n'):
            text_block = text_block.strip()
            if not text_block:
                continue
            new_p = doc.add_paragraph(style=style)
            if alignment is not None:
                new_p.alignment = alignment
            run = new_p.add_run(text_block)
            run.font.name = font_name
            if font_size:
                run.font.size = font_size

    doc.save(new_cl_path)

def convert_to_pdf_libreoffice(docx_path):
    pdf_dir = os.path.dirname(docx_path)
    soffice_path = os.environ.get('SOFFICE_PATH', '/Applications/LibreOffice.app/Contents/MacOS/soffice')
    subprocess.run([
        soffice_path,
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', pdf_dir,
        docx_path
    ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

async def generate_for_job(job_id, custom_instructions=None):
    import db
    conn = db.get_connection()
    cursor = conn.cursor()
    cursor.execute('''SELECT job_id, title, company, location, description, link, score, reasoning, status,
                             created_at, is_promoted, estimated_salary, is_recruiter, hiring_manager_name,
                             jd_language, application_notes
                      FROM jobs WHERE job_id = ?''', (job_id,))
    job = cursor.fetchone()
    if not job:
        print(f"Job {job_id} not found.")
        conn.close()
        return False
        
    groq_client = get_groq_client()
    gemini_client = get_gemini_client()
    
    if not groq_client and not gemini_client:
        print("Error: You must set either GROQ_API_KEY or GEMINI_API_KEY in your .env file.")
        conn.close()
        return False
        
    cv_text = extract_text(CV_PATH)
    try:
        with open(DOSSIER_PATH, 'r', encoding='utf-8') as f:
            dossier_text = f.read()
    except Exception as e:
        print(f"Error reading Career Dossier: {e}")
        conn.close()
        return False
        
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    job_id, title, company, location, description, link, score, reasoning, status, created_at, is_promoted, estimated_salary, is_recruiter, hiring_manager_name, jd_language, application_notes = job
    print(f"Generating documents for {company}: {title}...")
    
    baseline_pdf_path = CV_PATH.replace('.docx', '.pdf')
    if not os.path.exists(baseline_pdf_path):
        print("Converting baseline CV to PDF for page counting...")
        convert_to_pdf_libreoffice(CV_PATH)
    baseline_page_count = get_pdf_page_count(baseline_pdf_path) or 3
    
    max_attempts = 2
    for attempt in range(1, max_attempts + 1):
        try:
            texts = generate_tailored_texts(groq_client, gemini_client, job, cv_text, dossier_text, custom_instructions)
            if not texts:
                print(f"AI returned empty result for job {job_id}")
                cursor.execute('UPDATE jobs SET status = "failed" WHERE job_id = ?', (job_id,))
                conn.commit()
                conn.close()
                return False
        except Exception as e:
            print(f"Error generating texts for job {job_id}: {e}")
            cursor.execute('UPDATE jobs SET status = "failed" WHERE job_id = ?', (job_id,))
            conn.commit()
            conn.close()
            return False
            
        today = datetime.datetime.now().strftime("%Y%m%d")
        safe_company = "".join(x for x in company if x.isalnum())
        job_dir = os.path.join(OUTPUT_DIR, f"{today}_{safe_company}_{job_id}")
        
        import shutil
        if os.path.exists(OUTPUT_DIR):
            for d in os.listdir(OUTPUT_DIR):
                if d.endswith(f"_{job_id}"):
                    try:
                        shutil.rmtree(os.path.join(OUTPUT_DIR, d))
                    except Exception as ex:
                        print(f"Could not delete old directory {d}: {ex}")
                        
        os.makedirs(job_dir, exist_ok=True)
        
        new_cv_docx = os.path.join(job_dir, f"{today}_{safe_company}_CV.docx")
        new_cl_docx = os.path.join(job_dir, f"{today}_{safe_company}_CoverLetter.docx")
        new_cv_md = os.path.join(job_dir, f"{today}_{safe_company}_CV.md")
        new_cl_md = os.path.join(job_dir, f"{today}_{safe_company}_CoverLetter.md")
        
        # Adapt DOCX files
        adapt_cv(CV_PATH, new_cv_docx, texts)
        sanitize_generated_texts(texts, company, location, jd_language)
        company_hq = texts['company_hq']
        display_company = texts['display_company']
        adapt_cl(CL_PATH, new_cl_docx, texts['cover_letter_body'], company, display_company, company_hq, hiring_manager_name, jd_language)
        
        # Save MD format
        with open(new_cv_md, 'w') as f:
            f.write(extract_text(new_cv_docx))
        with open(new_cl_md, 'w') as f:
            f.write(extract_text(new_cl_docx))
        
        # Convert to PDF
        print(f"Converting DOCX to PDF for {company} (Attempt {attempt})...")
        try:
            convert_to_pdf_libreoffice(new_cv_docx)
            convert_to_pdf_libreoffice(new_cl_docx)
        except Exception as e:
            print(f"PDF conversion failed: {e}")
            
        new_cv_pdf = new_cv_docx.replace('.docx', '.pdf')
        new_page_count = get_pdf_page_count(new_cv_pdf)
        
        if new_page_count and new_page_count > baseline_page_count:
            if attempt < max_attempts:
                print(f"Generated CV is {new_page_count} pages (baseline is {baseline_page_count}). Retrying...")
                custom_instructions = (custom_instructions or "") + "\n\nCRITICAL: The previous generation was too long. You MUST shorten your bullet point replacements slightly to ensure the CV fits within its original page limit."
                continue
            else:
                print(f"Generated CV is still too long after {max_attempts} attempts. Keeping it.")
                break
        else:
            break
            
        
    # Save Job Description as PDF using Playwright
    print(f"Saving JD as PDF for {company}...")
    pdf_dir = os.path.join(CVS_DIR, 'jobs')
    os.makedirs(pdf_dir, exist_ok=True)
    safe_title = "".join(x for x in title if x.isalnum() or x in " -_")
    jd_pdf_filename = f"{job_id}_{safe_company}_{safe_title}.pdf"
    jd_pdf_path = os.path.join(pdf_dir, jd_pdf_filename)
    
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            html_content = f"<h1>{title} at {company}</h1><p>{location}</p><hr><pre style='white-space: pre-wrap;'>{description}</pre>"
            await page.set_content(html_content)
            await page.pdf(path=jd_pdf_path)
            await browser.close()
    except Exception as pdf_err:
        print(f"Failed to generate JD PDF for {job_id}: {pdf_err}")
        
    # Update status
    cursor.execute('UPDATE jobs SET status = "generated" WHERE job_id = ?', (job_id,))
    conn.commit()
    conn.close()
    print(f"Finished generating assets for {company}.")
    return True

async def run_generator():
    import db
    conn = db.get_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT job_id, title, company, description, reasoning, location FROM jobs WHERE status = "to_apply"')
    unprocessed_jobs = cursor.fetchall()
    conn.close()
    
    if not unprocessed_jobs:
        print("No jobs ready for document generation.")
        return
        
    total_jobs = len(unprocessed_jobs)
    for idx, job in enumerate(unprocessed_jobs, 1):
        if progress_tracker.is_stop_requested():
            print("Stop requested during generation!")
            break
        progress_tracker.set_status("Generating Assets", idx, total_jobs)
        
        job_id = job[0]
        print(f"[{idx}/{total_jobs}] Processing job {job_id}...")
        await generate_for_job(job_id)

if __name__ == '__main__':
    import asyncio
    asyncio.run(run_generator())
